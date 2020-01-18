import os
import csv
import json
from typing import List, Dict
import xml.etree.ElementTree as xml
from xml.dom import minidom
from wsgiref.util import FileWrapper

from django.utils import timezone
import docx
import docx.shared
import odf
import odf.table
from odf import teletype
import pandas as pd
import pylatex
import pylatex.utils

from musicavis.settings import EXPORTS_DIR
from app.models.practice import Practice
from app.backend.utils.enums import NewLine, FileType


class ExportPractices:
    __slots__ = ['username', 'practices', 'newline', 'fname', 'odfstyles']

    def __init__(self, username: str, practices: List[Practice], newline: NewLine):
        self.username = username
        self.practices = practices
        self.newline = newline.value
        self.fname = f'{EXPORTS_DIR}/{username}_practices'
        self.odfstyles = self.__prepare_odfstyles()

    def __prepare_odfstyles(self) -> Dict:
        h1style = odf.style.Style(name=f'Heading 1', family='paragraph')
        h1style.addElement(odf.style.ParagraphProperties(attributes={'textalign': 'center'}))
        h1style.addElement(odf.style.TextProperties(attributes={'fontsize': '18pt', 'fontweight': 'bold'}))

        boldstyle = odf.style.Style(name=f'Bold', family='text')
        boldprop = odf.style.TextProperties(fontweight='bold')
        boldstyle.addElement(boldprop)

        bulletliststyle = odf.text.ListStyle(name='BulletList')
        bulletlistproperty = odf.text.ListLevelStyleBullet(level='1', bulletchar='•')
        bulletlistproperty.addElement(odf.style.ListLevelProperties(minlabelwidth='1fcm'))
        bulletliststyle.addElement(bulletlistproperty)

        tablecontents = odf.style.Style(name='Normal Cells', family='table-cell')
        tablecontents.addElement(odf.style.TextProperties(fontfamily='Arial', fontsize='14pt'))

        widewidth = odf.style.Style(name='col1', family='table-column')
        widewidth.addElement(odf.style.TableColumnProperties(columnwidth='2.8cm', breakbefore='auto'))

        widewidth2 = odf.style.Style(name='col2', family='table-column')
        widewidth2.addElement(odf.style.TableColumnProperties(columnwidth='6.8cm', breakbefore='auto'))

        borderbottom = odf.style.Style(name='cell1', family='table-cell')
        borderbottom.addElement(odf.style.TableCellProperties(borderbottom='0.05pt solid #000000'))

        borderright = odf.style.Style(name='cell2', family='table-cell')
        borderright.addElement(odf.style.TableCellProperties(borderright='0.05pt solid #000000'))

        return {
            'h1': h1style,
            'bold': boldstyle,
            'bulletlist': bulletliststyle,
            'tablecontents': tablecontents,
            'widewidth': widewidth,
            'widewidth2': widewidth2,
            'borderbottom': borderbottom,
            'borderright': borderright,
        }

    def export(self, filetype: FileType) -> str:
        fname = f'{self.fname}_{timezone.now():%d%m%y}.{filetype.value}'
        if filetype == FileType.CSV:
            self.__export_csv(fname)
        elif filetype == FileType.DOCX:
            self.__export_docx(fname)
        elif filetype == FileType.XLSX:
            self.__export_excel(fname)
        elif filetype == FileType.XML:
            self.__export_xml(fname)
        elif filetype == FileType.TXT:
            self.__export_txt(fname)
        elif filetype == FileType.PDF:
            self.__export_pdf(fname)
        elif filetype == FileType.JSON:
            self.__export_json(fname)
        elif filetype == FileType.ODS:
            self.__export_ods(fname)
        elif filetype == FileType.ODT:
            self.__export_odt(fname)
        return os.path.basename(fname)

    def __export_csv(self, fname: str):
        with open(fname, 'w', newline=self.newline) as f:
            writer = csv.writer(f, delimiter='\n', quoting=csv.QUOTE_MINIMAL)
            for practice in self.practices:
                header = f'{practice.date},{practice.instrument.name.title()}'
                goals = 'goals: ' + ','.join([x.name for x in practice.goals.all()])
                exercises = 'exercises: ' + f',{self.newline}\t\t\t'.join([(f'{x.name},{x.bpm_start}bpm,{x.bpm_end}'
                                                                      f'bpm,{x.minutes}m')
                                                                     for x in practice.exercises.all()])
                improvements = 'improvements: ' + ','.join([x.name for x in practice.improvements.all()])
                positives = 'positives: ' + ','.join([x.name for x in practice.positives.all()])

                writer.writerow(([header, goals] + exercises.split('\n') +
                                 [improvements, positives, f'notes: {practice.notes}', '']))

    def __export_docx(self, fname: str):
        bullet = 'List Bullet'

        with open(fname, 'w', newline=self.newline):
            document = docx.Document()
            document.add_heading(f'Practice Sessions for {self.username.title()}', 0)

            for practice in self.practices:
                document.add_heading(f'{practice.instrument.name.title()} ({practice.date:%a, %B %m %Y})', 1)
                document.add_heading('Goals', 2)
                [document.add_paragraph(x.name, style=bullet) for x in practice.goals.all()]

                document.add_heading('Exercises', 2)
                [document.add_paragraph(f'{x.name} at {x.bpm_start}-{x.bpm_end}bpm for {x.minutes}m', bullet)
                 for x in practice.exercises.all()]

                document.add_heading('Improvements', 2)
                [document.add_paragraph(x.name, style=bullet) for x in practice.improvements.all()]

                document.add_heading('Positives', 2)
                [document.add_paragraph(x.name, style=bullet) for x in practice.positives.all()]

                document.add_heading('Notes', 2)
                document.add_paragraph('No notes.' if not practice.notes else practice.notes)

            document.save(fname)

    def __export_excel(self, fname: str):
        writer = pd.ExcelWriter(fname, engine='xlsxwriter')
        sheet_name = 'Practices'

        rowoffset = 0
        for index, practice in enumerate(self.practices):
            dfs = [pd.DataFrame({'Goals': [x.name for x in practice.goals.all()]}),
                   pd.DataFrame({'Exercise Name': [x.name for x in practice.exercises.all()],
                                 'Bpm Range': [f'{x.bpm_start}-{x.bpm_end}' for x in practice.exercises.all()],
                                 'Time (minutes)': [x.minutes for x in practice.exercises.all()]}),
                   pd.DataFrame({'Improvements': [x.name for x in practice.improvements.all()]}),
                   pd.DataFrame({'Positives': [x.name for x in practice.positives.all()]}),
                   pd.DataFrame({'Notes': [practice.notes]})]

            df_combined = pd.concat(dfs, axis=1)
            offset = index + rowoffset
            df_combined.to_excel(writer, sheet_name=sheet_name, startrow=offset + 1)

            merge_format = writer.book.add_format({'bold': 1, 'align': 'center', 'valign': 'vcenter', 'border': 1})
            header = f'Practice #{index + 1} - {practice.instrument.name.title()} - {practice.date:%a, %B %m %Y}'
            writer.sheets[sheet_name].merge_range(f'A{offset+1}:H{offset + 1}', header, merge_format)

            all_elements = [practice.goals.all(), practice.exercises.all(),
                            practice.improvements.all(), practice.positives.all()]
            rowoffset += len(max(all_elements, key=len)) + 5

        for i, width in enumerate([15, 30, 15, 15, 30, 30, 20]):
            writer.sheets[sheet_name].set_column(i + 1, i + 1, width)
        writer.save()

    def __export_xml(self, fname: str):
        practices = xml.Element('Practices')
        for index, practice in enumerate(self.practices):
            practice_root = xml.SubElement(practices, 'Practice')

            metadata = xml.SubElement(practice_root, 'Metadata')
            number = xml.SubElement(metadata, 'Number')
            number.text = str(index + 1)
            date = xml.SubElement(metadata, 'Date')
            date.text = str(practice.date)
            instrument = xml.SubElement(metadata, 'Instrument')
            instrument.text = practice.instrument.name.title()

            self.__create_name_nodes(practice_root, 'Goals', practice.goals.all())

            exercises_root = xml.SubElement(practice_root, 'Exercises')
            for x in practice.exercises.all():
                exercise = xml.SubElement(exercises_root, 'Exercise')
                name = xml.SubElement(exercise, 'Name')
                name.text = x.name
                bpm_range = xml.SubElement(exercise, 'BpmRange')
                bpm_range.text = f'{x.bpm_start}-{x.bpm_end}'
                minutes = xml.SubElement(exercise, 'Minutes')
                minutes.text = str(x.minutes)

            self.__create_name_nodes(practice_root, 'Improvements', practice.improvements.all())
            self.__create_name_nodes(practice_root, 'Positives', practice.positives.all())
            notes = xml.SubElement(practice_root, 'Notes')
            notes.text = practice.notes

        with open(fname, 'w', newline=self.newline) as f:
            raw = xml.tostring(practices, 'utf-8', method='xml')
            reparsed = minidom.parseString(raw)
            f.write(reparsed.toprettyxml(indent="  "))

    def __create_name_nodes(self, root: xml.Element, name: str, elements):
        sub = xml.SubElement(root, name)
        for el in elements:
            subelement = xml.SubElement(sub, 'Name')
            subelement.text = el.name

    def __export_txt(self, fname: str):
        with open(fname, 'w') as f:
            for practice in self.practices:
                header = f'{practice.instrument.name.title()} ({practice.date}){self.newline}'
                goals = f"goals -> {', '.join([x.name for x in practice.goals.all()])}{self.newline}"
                exercises = (f'exercises ->{self.newline}\t' +
                             f'{self.newline}\t'.join([f'{x.name} {x.bpm_start}-{x.bpm_end}bpm {x.minutes}m'
                                                       for x in practice.exercises.all()]) + f'{self.newline}')
                improvements = (f"improvements -> {', '.join([x.name for x in practice.improvements.all()])}"
                                f'{self.newline}')
                positives = f"positives -> {', '.join([x.name for x in practice.positives.all()])}{self.newline}"
                notes = f'notes -> {practice.notes}{self.newline}'
                f.writelines([header, goals, exercises, improvements, positives, notes, f'{self.newline}'])

    def __export_pdf(self, fname: str):
        geometry_options = {'margin': '1in'}
        doc = pylatex.Document(geometry_options=geometry_options)
        header = pylatex.PageStyle('header')

        with doc.create(pylatex.MiniPage(align='c')):
            doc.append(pylatex.HugeText(pylatex.utils.bold('Practice Sessions')))
            doc.append(pylatex.LineBreak())

        with header.create(pylatex.Foot('L')):
            header.append('@2019 Musicavis')

        with header.create(pylatex.Foot('R')):
            header.append(pylatex.simple_page_number())

        doc.preamble.append(header)
        doc.change_document_style('header')

        for practice in self.practices:
            with doc.create(pylatex.Section(f'{practice.instrument.name.title()} ({practice.date:%a, %B %m %Y})')):
                self.__add_list_itemize_latex(doc, 'Goals', practice.goals.all())

                with doc.create(pylatex.Subsection('Exercises')):
                    with doc.create(pylatex.Itemize()) as itemize:
                        for x in practice.exercises.all():
                            itemize.add_item(f'{x.name} at {x.bpm_start}-{x.bpm_end}bpm for {x.minutes}m')

                self.__add_list_itemize_latex(doc, 'Improvements', practice.improvements.all())
                self.__add_list_itemize_latex(doc, 'Positives', practice.positives.all())

            with doc.create(pylatex.Subsection('Notes:')):
                doc.append(practice.notes)

        doc.generate_pdf(fname[:-4], clean=True)

    def __add_list_itemize_latex(self, doc, title, elements):
        with doc.create(pylatex.Subsection(title)):
            with doc.create(pylatex.Itemize()) as itemize:
                [itemize.add_item(x.name) for x in elements]

    def __export_json(self, fname: str):
        root = {}
        for index, practice in enumerate(self.practices):
            root[f'practice-{index + 1}'] = {
                'metadata': {
                    'date': str(practice.date),
                    'instrument': practice.instrument.name.title()
                },
                'goals': [x.name for x in practice.goals.all()],
                'exercises': [{'name': x.name, 'bpm_range': f'{x.bpm_start}-{x.bpm_end}', 'minutes': float(x.minutes)}
                              for x in practice.exercises.all()],
                'improvements': [x.name for x in practice.improvements.all()],
                'positives': [x.name for x in practice.positives.all()],
                'notes': practice.notes
            }

        with open(fname, 'w') as f:
            f.write(json.dumps(root, indent=2))

    def __export_odt(self, fname: str):
        doc = odf.opendocument.OpenDocumentText()
        doc.styles.addElement(self.odfstyles['h1'])
        doc.automaticstyles.addElement(self.odfstyles['bold'])

        for index, practice in enumerate(self.practices):
            header = odf.text.H(outlinelevel=1, stylename=self.odfstyles['h1'],
                                text=(f'[{practice.date:%a, %d %B %Y}] Practice #{index + 1} - '
                                      f'{practice.instrument.name.title()}'))
            doc.text.addElement(header)
            doc.text.addElement(odf.text.P())

            self.__add_bullet_list_odf(doc, 'Goals:', [x.name for x in practice.goals.all()])

            exercises_para = self.__bold('Exercises:')
            bulletlist = odf.text.List(stylename=self.odfstyles['bulletlist'])
            for exercise in practice.exercises.all():
                listitem = odf.text.ListItem()
                listitem_p = odf.text.P(text=(f'{exercise.name} - {exercise.bpm_start} to '
                                              f'{exercise.bpm_end} bpm - {exercise.minutes}m'))
                listitem.addElement(listitem_p)
                bulletlist.addElement(listitem)
            doc.text.addElement(exercises_para)
            doc.text.addElement(bulletlist)
            doc.text.addElement(odf.text.P())

            self.__add_bullet_list_odf(doc, 'Improvements:', [x.name for x in practice.improvements.all()])
            self.__add_bullet_list_odf(doc, 'Positives:', [x.name for x in practice.positives.all()])

            notes = self.__bold('Notes: ')
            teletype.addTextToElement(notes, practice.notes if practice.notes is not None else '')
            doc.text.addElement(notes)
            doc.text.addElement(odf.text.P())
            doc.text.addElement(odf.text.P())

        doc.save(fname)

    def __add_bullet_list_odf(self, doc, title, elements):
        para = self.__bold(text=title)
        bulletlist = odf.text.List(stylename=self.odfstyles['bulletlist'])
        for el in elements:
            listitem = odf.text.ListItem()
            listitem_p = odf.text.P(text=el)
            listitem.addElement(listitem_p)
            bulletlist.addElement(listitem)
        doc.text.addElement(para)
        doc.text.addElement(bulletlist)
        doc.text.addElement(odf.text.P())

    def __export_ods(self, fname: str):
        doc = odf.opendocument.OpenDocumentSpreadsheet()
        doc.styles.addElement(self.odfstyles['tablecontents'])
        doc.automaticstyles.addElement(self.odfstyles['widewidth'])
        doc.automaticstyles.addElement(self.odfstyles['widewidth2'])
        doc.automaticstyles.addElement(self.odfstyles['borderbottom'])
        doc.automaticstyles.addElement(self.odfstyles['borderright'])
        doc.automaticstyles.addElement(self.odfstyles['bold'])

        table = odf.table.Table(name='Practices')
        table.addElement(odf.table.TableColumn(stylename=self.odfstyles['widewidth'], defaultcellstylename='ce1'))
        table.addElement(odf.table.TableColumn(stylename=self.odfstyles['widewidth2'], defaultcellstylename='ce2'))
        table.addElement(odf.table.TableColumn(stylename=self.odfstyles['widewidth2'], defaultcellstylename='ce3'))
        table.addElement(odf.table.TableColumn(stylename=self.odfstyles['widewidth2'], defaultcellstylename='ce4'))

        for index, practice in enumerate(self.practices):
            tr_header = self.__add_row(table)
            self.__add_string_cell(tr_header, '', border=self.odfstyles['borderbottom'])
            self.__add_string_cell(tr_header, f'Practice #{index +1}', True, self.odfstyles['borderbottom'])
            self.__add_string_cell(tr_header, practice.instrument.name.title(), True, self.odfstyles['borderbottom'])
            self.__add_string_cell(tr_header, f'{practice.date:%a, %d %B %Y}', True, self.odfstyles['borderbottom'])

            tr_goals = self.__add_row(table)
            self.__add_string_cell(tr_goals, 'Goals:', True, self.odfstyles['borderright'])
            [self.__add_string_cell(tr_goals, goal.name) for goal in practice.goals.all()]

            tr_exercise = self.__add_row(table)
            self.__add_string_cell(tr_exercise, 'Exercises:', True, self.odfstyles['borderright'])
            exercises = practice.exercises.all()
            for index, exercise in enumerate(exercises):
                self.__add_string_cell(tr_exercise, exercise.name)
                self.__add_string_cell(tr_exercise, f'{exercise.bpm_start}-{exercise.bpm_end} bpm')
                self.__add_string_cell(tr_exercise, f'{exercise.minutes}m')

                if index + 1 != len(exercises):
                    tr_exercise = self.__add_row(table)
                    self.__add_string_cell(tr_exercise, odf.text.P(), border=self.odfstyles['borderright'])

            tr_improvements = self.__add_row(table)
            self.__add_string_cell(tr_improvements, 'Improvements:', True, self.odfstyles['borderright'])
            [self.__add_string_cell(tr_improvements, improvement.name) for improvement in practice.improvements.all()]

            tr_positives = self.__add_row(table)
            self.__add_string_cell(tr_positives, 'Positives:', True, self.odfstyles['borderright'])
            [self.__add_string_cell(tr_positives, positive.name) for positive in practice.positives.all()]

            tr_notes = self.__add_row(table)
            self.__add_string_cell(tr_notes, f'Notes:', True, self.odfstyles['borderright'])
            self.__add_string_cell(tr_notes, practice.notes)
            self.__add_row(table)

        doc.spreadsheet.addElement(table)
        doc.save(fname)

    def __add_row(self, table):
        tr = odf.table.TableRow()
        table.addElement(tr)
        return tr

    def __add_string_cell(self, row, text, bold=False, border=None):
        cell = odf.table.TableCell(valuetype='string', stylename=border)

        if bold:
            p = self.__bold(text=text)
            cell.addElement(p)
        else:
            cell.addElement(odf.text.P(text=text))

        row.addElement(cell)

    def __bold(self, text):
        p = odf.text.P()
        boldpart = odf.text.Span(stylename=self.odfstyles['bold'], text=text)
        p.addElement(boldpart)
        return p


class FileDeleteWrapper(FileWrapper):

    def __init__(self, filepath, *args, **kwargs):
        self.filepath = filepath
        super(FileDeleteWrapper, self).__init__(*args, **kwargs)

    def __del__(self, *args, **kwargs):
        os.remove(self.filepath)
